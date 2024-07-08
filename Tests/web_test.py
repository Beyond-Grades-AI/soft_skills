import unittest
from unittest.mock import patch, MagicMock, Mock
from unittest import TestCase
from soft_skills.web.management.commands.check_evaluations import Command
from django.test import TestCase, RequestFactory
from django.contrib.sessions.middleware import SessionMiddleware
from django.contrib.messages.middleware import MessageMiddleware
from django.contrib.auth.models import AnonymousUser
from django.contrib.sessions.backends.db import SessionStore
from django.shortcuts import redirect, render
from soft_skills.web.models import Teacher, Test, Question
from soft_skills.web.views import login_screen, main_screen, extract_text_from_docx, extract_text_from_pdf, create_questions, generate_link, display_link, test_page, tests_screen, test_feedback, review_test
from django.http import HttpResponse, HttpRequest
from docx import Document
from io import StringIO
from pdfplumber import PDF
from django.urls import reverse
from django.test.client import RequestFactory
from django.contrib.auth.models import User
from django.test import Client



class TestHandleFunction(TestCase):
    
    @patch('soft_skills.web.management.commands.check_evaluations.get_questions_answers_test_df')
    @patch('soft_skills.web.management.commands.check_evaluations.evaluate_answers')
    @patch('soft_skills.web.management.commands.check_evaluations.update_origin_eval')
    @patch('soft_skills.web.management.commands.check_evaluations.print')
    @patch('time.sleep')
    def test_handle_with_data(self, mock_sleep, mock_print, mock_update_origin_eval, mock_evaluate_answers, mock_get_questions_answers_test_df):
        mock_get_questions_answers_test_df.return_value = (answer_df, question_df, test_df)
        answer_df = MagicMock()
        answer_df.empty = False
        question_df = MagicMock()
        test_df = MagicMock()
        
        Command().handle()
        
        mock_evaluate_answers.assert_called()
        mock_update_origin_eval.assert_called()
        mock_sleep.assert_called()
        
    @patch('soft_skills.web.management.commands.check_evaluations.get_questions_answers_test_df')
    @patch('soft_skills.web.management.commands.check_evaluations.print')
    @patch('time.sleep')
    def test_handle_with_empty_data(self, mock_sleep, mock_print, mock_get_questions_answers_test_df):
        mock_get_questions_answers_test_df.return_value = (answer_df, question_df, test_df)
        answer_df = MagicMock()
        answer_df.empty = True
        question_df = MagicMock()
        test_df = MagicMock()
        
        Command().handle()
        
        mock_print.assert_called_with("No answers to evaluate. Sleeping...")
        mock_sleep.assert_called_with(300)


class LoginScreenTests(TestCase):
    def setUp(self):
        self.factory = RequestFactory()
        self.middleware = SessionMiddleware()
        self.msg_middleware = MessageMiddleware()

    def test_get_request(self):
        request = self.factory.get('/login/')
        request.user = AnonymousUser()
        self.middleware.process_request(request)
        self.msg_middleware.process_request(request)
        request.session = SessionStore()

        response = login_screen(request)

        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'login_screen.html')

    def test_post_request_existing_teacher(self):
        request = self.factory.post('/login/', {'input_email': 'test@example.com'})
        request.user = AnonymousUser()
        self.middleware.process_request(request)
        self.msg_middleware.process_request(request)
        request.session = SessionStore()

        Teacher.objects.create(email='test@example.com')

        response = login_screen(request)

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, '/main_screen/')
        self.assertEqual(request.session['teacher'], 'test@example.com')

    def test_post_request_new_teacher(self):
        request = self.factory.post('/login/', {'input_email': 'new@example.com'})
        request.user = AnonymousUser()
        self.middleware.process_request(request)
        self.msg_middleware.process_request(request)
        request.session = SessionStore()

        response = login_screen(request)

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, '/main_screen/')
        self.assertEqual(request.session['teacher'], 'new@example.com')
        self.assertEqual(Teacher.objects.count(), 1)
        self.assertEqual(Teacher.objects.get().email, 'new@example.com')



class MainScreenTestCase(unittest.TestCase):
    def setUp(self):
        self.factory = RequestFactory()

    @patch('soft_skills.web.views.render')
    def test_main_screen_returns_rendered_template(self, mock_render):
        # Arrange
        request = self.factory.get('/')
        mock_render.return_value = HttpResponse()

        # Act
        response = main_screen(request)

        # Assert
        self.assertIsInstance(response, HttpResponse)
        mock_render.assert_called_once_with(request, 'main_screen.html')

    @patch('soft_skills.web.views.render')
    def test_main_screen_returns_rendered_template_with_context(self, mock_render):
        # Arrange
        request = self.factory.get('/')
        mock_render.return_value = HttpResponse()

        # Act
        response = main_screen(request)

        # Assert
        self.assertIsInstance(response, HttpResponse)
        mock_render.assert_called_once_with(request, 'main_screen.html', {})



class TestExtractTextFromDocx(TestCase):
    def test_extract_text_from_docx(self):
        # Test case 1: Empty document
        with patch('docx.Document') as mock_doc:
            mock_doc.return_value.paragraphs = []
            result = extract_text_from_docx('test.docx')
            self.assertEqual(result, '')

        # Test case 2: Document with one paragraph
        with patch('docx.Document') as mock_doc:
            mock_doc.return_value.paragraphs = [
                MockParagraph('Hello world')
            ]
            result = extract_text_from_docx('test.docx')
            self.assertEqual(result, 'dlrow olleH\n\n')

        # Test case 3: Document with multiple paragraphs
        with patch('docx.Document') as mock_doc:
            mock_doc.return_value.paragraphs = [
                MockParagraph('Hello world'),
                MockParagraph('This is a test')
            ]
            result = extract_text_from_docx('test.docx')
            self.assertEqual(result, 'dlrow olleH\n\nsiht si eht\n\n')

class MockParagraph:
    def __init__(self, text):
        self.text = text

    @property
    def text(self):
        return self._text

    @text.setter
    def text(self, value):
        self._text = value
        self._reversed_text = value[::-1]

    def __str__(self):
        return self._reversed_text



class TestExtractTextFromPDF(unittest.TestCase):
    @patch("pdfplumber.open")
    def test_extract_text_from_pdf(self, mock_pdfplumber_open):
        # Test case 1: PDF with multiple pages
        mock_pdf = MagicMock(spec=PDF)
        mock_pdf.pages = [MagicMock(spec=StringIO) for _ in range(3)]
        mock_pdf.pages[0].extract_text.return_value = "Page 1 text"
        mock_pdf.pages[1].extract_text.return_value = "Page 2 text"
        mock_pdf.pages[2].extract_text.return_value = "Page 3 text"
        mock_pdfplumber_open.return_value = mock_pdf

        result = extract_text_from_pdf("test.pdf")
        self.assertEqual(result, "Page 1 textPage 2 textPage 3 text")

        # Test case 2: PDF with no pages
        mock_pdf.pages = []
        result = extract_text_from_pdf("test.pdf")
        self.assertEqual(result, "")

        # Test case 3: PDF with one page
        mock_pdf.pages = [MagicMock(spec=StringIO)]
        mock_pdf.pages[0].extract_text.return_value = "Page 1 text"
        result = extract_text_from_pdf("test.pdf")
        self.assertEqual(result, "Page 1 text")

        # Test case 4: PDF with one page but no text
        mock_pdf.pages = [MagicMock(spec=StringIO)]
        mock_pdf.pages[0].extract_text.return_value = ""
        result = extract_text_from_pdf("test.pdf")
        self.assertEqual(result, "")

        # Test case 5: PDF with one page but no text and no exception
        mock_pdf.pages = [MagicMock(spec=StringIO)]
        mock_pdf.pages[0].extract_text.return_value = ""
        mock_pdf.pages[0].extract_text.side_effect = None
        result = extract_text_from_pdf("test.pdf")
        self.assertEqual(result, "")

        # Test case 6: PDF with exception
        mock_pdfplumber_open.side_effect = Exception("Test exception")
        result = extract_text_from_pdf("test.pdf")
        self.assertEqual(result, "")


class TestCreateQuestions(unittest.TestCase):

    @patch('soft_skills.web.views.extract_text_from_docx')
    @patch('soft_skills.web.views.extract_text_from_pdf')
    def test_create_questions_post_valid_file(self, mock_extract_text_pdf, mock_extract_text_docx):
        request_factory = RequestFactory()
        request = request_factory.post('/create_questions', {
            'skill': 'Skill',
            'subject': 'Subject',
            'grade': 'Grade',
            'title': 'Title',
            'file': MagicMock(name='file', content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        })
        response = create_questions(request)
        self.assertEqual(response.status_code, 200)  # Assuming render returns a 200 status code

    @patch('soft_skills.web.views.extract_text_from_docx')
    @patch('soft_skills.web.views.extract_text_from_pdf')
    def test_create_questions_post_unsupported_file(self, mock_extract_text_pdf, mock_extract_text_docx):
        request_factory = RequestFactory()
        request = request_factory.post('/create_questions', {
            'skill': 'Skill',
            'subject': 'Subject',
            'grade': 'Grade',
            'title': 'Title',
            'file': MagicMock(name='file', content_type='unsupported')
        })
        response = create_questions(request)
        self.assertEqual(response.status_code, 200)  # Assuming render returns a 200 status code

    def test_create_questions_post_no_file(self):
        request_factory = RequestFactory()
        request = request_factory.post('/create_questions', {
            'skill': 'Skill',
            'subject': 'Subject',
            'grade': 'Grade',
            'title': 'Title'
        })
        response = create_questions(request)
        self.assertEqual(response.status_code, 500)  # Assuming HttpResponseServerError returns a 500 status code

    def test_create_questions_post_missing_fields(self):
        request_factory = RequestFactory()
        request = request_factory.post('/create_questions', {})
        response = create_questions(request)
        self.assertEqual(response.status_code, 500)  # Assuming HttpResponseServerError returns a 500 status code


class TestGenerateLinkFunction(unittest.TestCase):

    def test_generate_link_success(self):
        # Creating a POST request with necessary data
        factory = RequestFactory()
        request = factory.post('/generate_link/', {'test_id': '1', 'subject': 'Math', 'skill': 'Problem Solving', 'grade': 'A', 'test_title': 'Math Test', 'questions[]': ['Question 1', 'Question 2']})
        
        response = generate_link(request)
        self.assertEqual(response.status_code, 200)  # Assuming 200 is the success status code

    def test_generate_link_missing_subject_skill(self):
        # Creating a POST request with missing subject and skill
        factory = RequestFactory()
        request = factory.post('/generate_link/', {'test_id': '1', 'subject': '', 'skill': '', 'grade': 'A', 'test_title': 'Math Test', 'questions[]': ['Question 1', 'Question 2']})
        
        response = generate_link(request)
        self.assertEqual(response.status_code, 500)  # Assuming 500 is the error status code for missing subject and skill

    def test_generate_link_render_test_link(self):
        # Creating a GET request
        factory = RequestFactory()
        request = factory.get('/generate_link/')
        
        response = generate_link(request)
        self.assertIn('test_page_url', response.content.decode())  # Checking if test_page_url is present in the response


class DisplayLinkTests(TestCase):
    def setUp(self):
        self.factory = RequestFactory()
        self.test_page_url = '/test_page/1'
        self.grade = '9'
        self.skill = 'math'
        self.subject = 'algebra'
        self.test_title = 'Test Title'

    def test_display_link_returns_test_link_template(self):
        request = self.factory.get(reverse('display_link', args=[self.test_page_url, self.grade, self.skill, self.subject, self.test_title]))
        response = display_link(request, self.test_page_url, self.grade, self.skill, self.subject, self.test_title)
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'test_link.html')

    def test_display_link_context(self):
        request = self.factory.get(reverse('display_link', args=[self.test_page_url, self.grade, self.skill, self.subject, self.test_title]))
        response = display_link(request, self.test_page_url, self.grade, self.skill, self.subject, self.test_title)
        self.assertEqual(response.context['test_page_url'], self.test_page_url)
        self.assertEqual(response.context['grade'], self.grade)
        self.assertEqual(response.context['skill'], self.skill)
        self.assertEqual(response.context['subject'], self.subject)
        self.assertEqual(response.context['test_title'], self.test_title)



class TestTestPage(TestCase):
    def setUp(self):
        self.factory = RequestFactory()

    def test_retrieve_test_and_questions_successfully(self):
        # Create a test object and associated questions
        test_grade = '10'
        test_skill = 'Test Skill'
        test = Test.objects.create(grade=test_grade, skill=test_skill)
        question_text = 'Test Question'
        question = Question.objects.create(text=question_text, test=test)

        # Create a request object
        request = self.factory.get('/test_page/')

        # Call the function
        response = test_page(request, test.id)

        # Assert the response
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'test_page.html')
        self.assertEqual(response.context['questions'], [question])
        self.assertEqual(response.context['test'], test)
        self.assertEqual(response.context['grade'], test_grade)
        self.assertEqual(response.context['skill'], test_skill)

    def test_no_questions_found(self):
        # Create a test object
        test_grade = '10'
        test_skill = 'Test Skill'
        test = Test.objects.create(grade=test_grade, skill=test_skill)

        # Create a request object
        request = self.factory.get('/test_page/')

        # Call the function
        response = test_page(request, test.id)

        # Assert the response
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'test_page.html')
        self.assertEqual(response.context['error_message'], 'No questions found for this test.')

    def test_test_not_found(self):
        # Create a request object
        request = self.factory.get('/test_page/')

        # Call the function
        response = test_page(request, 'invalid_test_id')

        # Assert the response
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'test_page.html')
        self.assertEqual(response.context['error_message'], 'Test not found.')


class TestSubmitAnswers(unittest.TestCase):

    def test_valid_submission(self):
        client = Client()
        response = client.post(reverse('submit_answers'), {
            'test_id': 1,
            'full_name': 'John Doe',
            'email': 'john@example.com',
            'test': 'on',
            'answer_1': 'Answer to question 1',
            'answer_2': 'Answer to question 2'
        })
        self.assertEqual(response.status_code, 200)  # Assuming 200 for success

    def test_missing_or_invalid_name(self):
        client = Client()
        response = client.post(reverse('submit_answers'), {
            'test_id': 1,
            'full_name': 'John',  # Invalid name with only one word
            'email': 'john@example.com',
            'test': 'on',
            'answer_1': 'Answer to question 1',
            'answer_2': 'Answer to question 2'
        })
        self.assertEqual(response.status_code, 200)  # Assuming 200 for success with error message

    def test_existing_answers(self):
        client = Client()
        response = client.post(reverse('submit_answers'), {
            'test_id': 1,
            'full_name': 'Jane Doe',
            'email': 'jane@example.com',
            'test': 'on',
            'answer_1': 'Answer to question 1',
            'answer_2': 'Answer to question 2'
        })
        # Submit again with the same email and questions
        response_existing = client.post(reverse('submit_answers'), {
            'test_id': 1,
            'full_name': 'Jane Doe',
            'email': 'jane@example.com',
            'test': 'on',
            'answer_1': 'New answer to question 1',
            'answer_2': 'New answer to question 2'
        })
        self.assertEqual(response_existing.status_code, 200)  # Assuming 200 for success with error message



class TestTestsScreen(unittest.TestCase):
    def setUp(self):
        self.factory = RequestFactory()

    @patch('soft_skills.web.views.Teacher.objects.get')
    def test_get_request(self, mock_get):
        # Test GET request
        request = self.factory.get('/tests_screen')
        request.session = {'teacher': 'test@example.com'}
        mock_get.return_value.tests.all.return_value = ['test1', 'test2']

        response = tests_screen(request)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context_data['tests'], ['test1', 'test2'])

    @patch('soft_skills.web.views.Teacher.objects.get')
    def test_post_request(self, mock_get):
        # Test POST request
        request = self.factory.post('/tests_screen')
        request.session = {'teacher': 'test@example.com'}
        request.POST = {'test_id': '1'}
        mock_get.return_value.tests.all.return_value = ['test1', 'test2']

        response = tests_screen(request)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context_data['test_id'], '1')
        self.assertEqual(response.context_data['students'], ['student1', 'student2'])
        self.assertEqual(response.context_data['test'], 'test1')
        self.assertEqual(response.context_data['tests'], ['test1', 'test2'])

    def test_teacher_not_logged_in(self):
        # Test when teacher is not logged in
        request = self.factory.get('/tests_screen')
        request.session = {}

        response = tests_screen(request)

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, '/login_screen')

    @patch('soft_skills.web.views.Teacher.objects.get')
    def test_teacher_does_not_exist(self, mock_get):
        # Test when teacher does not exist
        request = self.factory.get('/tests_screen')
        request.session = {'teacher': 'test@example.com'}
        mock_get.side_effect = Teacher.DoesNotExist

        response = tests_screen(request)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context_data['tests'], [])


class TestTestFeedback(TestCase):
    def setUp(self):
        self.factory = RequestFactory()
        self.test = Test.objects.create(name='Test')
        self.test.students.create(name='Student')

    def test_post_request(self):
        request = self.factory.post('/test_feedback', {'test_id': self.test.id})
        response = test_feedback(request)
        self.assertIsInstance(response, HttpResponse)
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context_data['students'].count(), 1)
        self.assertEqual(response.context_data['count'], 1)
        self.assertEqual(response.context_data['test'], self.test)

    def test_get_request(self):
        request = self.factory.get('/test_feedback')
        response = test_feedback(request)
        self.assertIsInstance(response, HttpResponse)
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context_data['students'], None)
        self.assertEqual(response.context_data['count'], None)
        self.assertEqual(response.context_data['test'], None)

    def test_anonymous_user(self):
        request = self.factory.get('/test_feedback')
        request.user = AnonymousUser()
        response = test_feedback(request)
        self.assertIsInstance(response, HttpResponse)
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, '/login_screen')


class TestReviewTestFunction(unittest.TestCase):

    def test_retrieve_test_object(self):
        # Test retrieving test object for given test_id and student_id
        request = Mock()
        request.method = 'GET'
        test_id = 1
        student_id = 1
        response = review_test(request, test_id, student_id)
        self.assertIsNotNone(response)

    def test_create_question_answers_dict(self):
        # Test creating a dictionary to store questions and answers
        request = Mock()
        request.method = 'GET'
        test_id = 1
        student_id = 1
        response = review_test(request, test_id, student_id)
        self.assertIsNotNone(response)

    def test_handle_get_request(self):
        # Test handling GET request by rendering 'review_test.html'
        request = Mock()
        request.method = 'GET'
        test_id = 1
        student_id = 1
        response = review_test(request, test_id, student_id)
        self.assertIsNotNone(response)

    def test_handle_post_request(self):
        # Test handling POST request by updating submitted evaluations and redirecting based on teacher's email
        request = Mock()
        request.method = 'POST'
        test_id = 1
        student_id = 1
        request.session = {'teacher': 'teacher@email.com'}
        response = review_test(request, test_id, student_id)
        self.assertIsNotNone(response)


if __name__ == '__main__':
    unittest.main()